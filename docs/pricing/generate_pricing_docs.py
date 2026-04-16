"""
Generate GrantAQ Pricing Strategy & Competitive Analysis documents.
Outputs: .docx (Word/Google Doc compatible) + .xlsx (Excel/Sheets compatible)
Updated with deep competitive intelligence research.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from datetime import datetime

TEAL = RGBColor(13, 148, 136)
DARK = RGBColor(28, 25, 23)
GRAY = RGBColor(120, 113, 108)

# ============================================================================
# WORD DOCUMENT
# ============================================================================

def create_word_doc():
    doc = Document()

    style = doc.styles['Heading 1']
    style.font.color.rgb = TEAL
    style.font.size = Pt(24)
    style2 = doc.styles['Heading 2']
    style2.font.color.rgb = DARK
    style2.font.size = Pt(18)
    style3 = doc.styles['Heading 3']
    style3.font.color.rgb = TEAL
    style3.font.size = Pt(14)

    doc.add_heading('GrantAQ — Pricing Strategy, Competitive Analysis & Growth Playbook', level=1)
    doc.add_paragraph(f'Prepared: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph('Confidential — Internal Use Only')
    doc.add_paragraph('')

    # ── MARKET OVERVIEW ──
    doc.add_heading('Market Overview', level=2)
    doc.add_paragraph(
        'The US charitable giving market reached $592.50B in 2024 (up 6.3% YoY). Foundation grantmaking '
        'exceeded $100B for the third consecutive year, with foundation assets hitting a record $1.6T. '
        'Federal grants represent 40% of government funding, with health sector comprising 60% of federal grants. '
        'The grant technology software market is valued at $2.2-2.66B (2024) and projected to reach $5.8-16.56B '
        'by 2034-2035, growing at 8.4-12.6% CAGR.'
    )
    doc.add_paragraph(
        'Key growth drivers: cloud-based adoption, AI-powered proposal writing, automation and compliance '
        'requirements, demand for data-driven decision making, and multi-funder portfolio management.'
    )

    # ── EXECUTIVE SUMMARY ──
    doc.add_heading('Executive Summary', level=2)
    doc.add_paragraph(
        'GrantAQ operates in the grant discovery, readiness assessment, and grant writing market. '
        'Our primary competitors are Instrumentl ($299-$899/mo, $55M raised, 4,500 customers), '
        'Submittable ($66.6M revenue, 11,000 customers), GrantWatch ($49/mo, 350K+ users), '
        'and Grantable ($50-$150/mo, 13,000 users). GrantAQ is uniquely positioned as the ONLY platform '
        'combining AI-powered grant matching, automated readiness diagnostics, eligibility assessment, '
        'and grant writing in a single product — at price points 50-90% below traditional consulting.'
    )

    # ── OUR SERVICES & RECOMMENDED PRICING ──
    doc.add_heading('Our Services & Recommended Pricing', level=2)

    doc.add_heading('Subscription Tiers (Recurring Revenue)', level=3)
    table = doc.add_table(rows=6, cols=6)
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(['Tier', 'Name', 'Monthly', 'Annual', 'Key Features', 'Target Customer']):
        table.rows[0].cells[i].text = h
    tiers = [
        ['Free', 'Explorer', '$0', 'Forever', '1 match/mo, federal only, 1 readiness, free eligibility check', 'Lead capture — everyone starts here'],
        ['Starter', 'Seeker', '$39/mo', '$390/yr', 'Full library, 10 pipeline, calendar, 5 docs, 15 chats', 'Small orgs exploring grants for the first time'],
        ['Pro', 'Strategist', '$99/mo', '$950/yr', 'Unlimited scorecard, vault, 30 chats, analytics, A-Z tracking', 'Active grant seekers applying to multiple programs'],
        ['Growth', 'Applicant', '$179/mo', '$1,710/yr', 'AI writing, expert review, unlimited pipeline, narratives', 'Orgs ready to write and submit applications'],
        ['Enterprise', 'Organization', '$349/mo', '$3,350/yr', 'Everything + dedicated writer, unlimited team, API, CSM', 'Multi-user orgs and consultants managing clients'],
    ]
    for r, row_data in enumerate(tiers, 1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val
    doc.add_paragraph('')

    doc.add_heading('One-Time Assessment & Service Tiers', level=3)
    table2 = doc.add_table(rows=7, cols=5)
    table2.style = 'Light Grid Accent 1'
    for i, h in enumerate(['Service', 'Price', 'Delivery', "What's Included", 'Why This Price']):
        table2.rows[0].cells[i].text = h
    services = [
        ['Grant Eligibility Status', 'FREE', 'Instant (AI)', 'Quick eligibility verdict, grant categories, blockers, quick wins', 'Lead magnet — costs $0.15/check, captures leads, funnels to diagnostic'],
        ['Readiness Diagnostic', 'FREE (with account)', '30 seconds (AI)', '10-step diagnostic: 5-layer audit, COSO controls, site-visit sim, funder matches, roadmap', 'Drives account creation. Every gap = upsell to Tier 1-3. Consultants charge $2,500+'],
        ['Tier 1 — Readiness Review', '$497', '5-7 days', 'Full diagnostic + 45-min walkthrough call', '80% cheaper than consultants ($2,500+). Below "need approval" threshold.'],
        ['Tier 2 — Remediation Roadmap', '$1,997', '2-3 weeks', 'Tier 1 + playbook, templates, vendor directory, 2 strategy calls, 30-day support', '60-87% below consultant rates. $25K grant win = 12x return.'],
        ['Tier 3 — Readiness Accelerator', '$4,997', '60-120 days', 'Tier 2 + SAM/UEI done-for-you, policies drafted, first application, weekly sessions', '50-75% below consultant rates. Target: $500K+ revenue orgs.'],
        ['Strategic Restructuring', '$1,497', '2-4 weeks', 'Structural fit analysis, restructuring options, alt capital roadmap, 60-min strategy call', 'Unique — no competitor has standalone restructuring product.'],
    ]
    for r, row_data in enumerate(services, 1):
        for c, val in enumerate(row_data):
            table2.rows[r].cells[c].text = val
    doc.add_paragraph('')

    doc.add_heading('Grant Writing Tiers (Per Grant)', level=3)
    table3 = doc.add_table(rows=5, cols=5)
    table3.style = 'Light Grid Accent 1'
    for i, h in enumerate(['Writing Tier', 'State/Foundation', 'Federal', 'SBIR/STTR', "What's Included"]):
        table3.rows[0].cells[i].text = h
    writing = [
        ['AI Only', '$149', '$349', '$499', 'AI-generated draft + compliance check'],
        ['AI + Audit', '$249', '$549', '$749', 'AI draft + human expert review + revision'],
        ['Expert-Led', '$499', '$1,249', '$1,749', 'Expert writer + AI tools + 2 revision rounds'],
        ['Full Confidence', '$0 upfront', '$0 upfront', '$0 upfront', 'Success fee: 3-10% of award. Win or pay nothing.'],
    ]
    for r, row_data in enumerate(writing, 1):
        for c, val in enumerate(row_data):
            table3.rows[r].cells[c].text = val
    doc.add_paragraph('')

    # ── DEEP COMPETITOR PROFILES ──
    doc.add_heading('Deep Competitor Profiles', level=2)

    competitors = [
        ('Instrumentl',
         'Founded: 2015, Oakland CA\n'
         'Funding: $55M+ raised (Series from Summit Partners, April 2025)\n'
         'Customers: 4,500+\n'
         'Pricing: $299/mo (Standard), $499/mo (Pro), $899/mo (Advanced)\n'
         'Database: 450,000+ funder profiles, 27,000+ active opportunities, 22,000+ RFPs\n'
         'Revenue: Private (customers report winning avg $1.1M more/year in grants)\n'
         'Data Sources: Grants.gov API, IRS 990 filings, state databases, corporate funders, manual research team\n'
         'Key Strength: Largest customer base, strong brand, big funding\n'
         'Key Weakness: No readiness assessment, no writing, expensive ($3,588-$10,788/year)\n'
         'What We Beat Them On: We have readiness diagnostics, eligibility checks, writing, and service tiers. They have none of these.'),

        ('Submittable',
         'Founded: Missoula, MT\n'
         'Revenue: $66.6M (2024), up from $49M (2023)\n'
         'Funding: $74.6M total (Series C: $47M, 2022; backed by True Ventures, Accel-KKR)\n'
         'Team: 230 employees, 52 sales reps, 48 engineers\n'
         'Customers: 11,000\n'
         'Business Model: Grant management SaaS + CSR platform (NOT discovery)\n'
         'Data Sources: Not a discovery tool — manages applications/submissions\n'
         'Key Strength: Revenue leader, diversified use cases\n'
         'Key Weakness: Not a direct competitor — they manage, we discover. Different market.\n'
         'What We Beat Them On: We serve grantseekers. They serve grantmakers and CSR programs.'),

        ('GrantWatch',
         'Founded: 2010 by Libby Hikind (29-year NYC teacher)\n'
         'Funding: Bootstrapped (no VC)\n'
         'Users: 350,000+ monthly active, 200,000-230,000 monthly visitors\n'
         'Pricing: $49/mo or $249/year\n'
         'Database: 10,000-11,500 active grants in 60+ categories, human-curated\n'
         'Data Sources: 40+ human researchers manually find and verify grants\n'
         'Key Strength: Largest user base, human curation = high quality\n'
         'Key Weakness: No AI matching, no readiness, no writing — just a searchable directory\n'
         'What We Beat Them On: AI matching, eligibility diagnostics, writing, automated nurture. They have none.'),

        ('GrantStation',
         'Founded: 1999 (acquired by Elios Media Group, September 2024)\n'
         'Funding: PE-backed (Elios Media)\n'
         'Pricing: $699/year or $1,782 for 3 years\n'
         'Database: 150,000+ funding profiles, 15,000+ curated grants\n'
         'Data Sources: 25+ year proprietary database, direct funder relationships, manual research\n'
         'Key Strength: Longevity (25+ years), expert credibility, deep funder relationships\n'
         'Key Weakness: Dated UI, no AI, no assessment tools, expensive for what you get\n'
         'What We Beat Them On: AI matching, instant diagnostics, modern UX, lower price point.'),

        ('Grantable',
         'Founded: 2020, Richmond VA\n'
         'Funding: $100K (grant-based, not VC — Virginia Innovation Partnership Corp)\n'
         'Users: 13,000+\n'
         'Pricing: $22/mo (Basic), $68/mo (Professional)\n'
         'Database: 130,000+ foundations from IRS 990 filings\n'
         'Data Sources: IRS 990 filings (ProPublica API), Grants.gov for federal, public records\n'
         'Key Strength: AI "coworker" that learns from past proposals, affordable\n'
         'Key Weakness: No readiness assessment, no eligibility diagnostics, no expert writing option\n'
         'What We Beat Them On: Readiness diagnostics, eligibility checks, expert writing tiers, service tiers.'),

        ('Granted AI',
         'Pricing: $57/mo (professional, annual billing)\n'
         'Database: 133,000+ foundations\n'
         'Business Model: Subscription with "Win a grant or money back" guarantee (12 months)\n'
         'Data Sources: Federal, state, foundation databases\n'
         'Key Strength: Performance guarantee, AI-first, unlimited drafts\n'
         'Key Weakness: Small database, no matching AI, no readiness tools\n'
         'What We Beat Them On: Matching AI, readiness diagnostics, eligibility checks, larger DB.'),

        ('Candid (Foundation Directory Online)',
         'Formed: 2019 merger of GuideStar + Foundation Center\n'
         'Business Model: Nonprofit org — 95% of users access data free. Revenue from licensing.\n'
         'Database: 325,000+ grantmakers, 2.4M+ recently awarded grants, $180B+ in funding tracked\n'
         'Data Sources: IRS 990s/990-PFs, IRS Exempt Org Master File, CA AG Registry, direct funder submissions\n'
         'Key Strength: THE authoritative data source. Powers 200+ giving platforms and DAFs.\n'
         'Key Weakness: Research tool only — no application help, no writing, no readiness\n'
         'What We Beat Them On: Everything except raw data volume. We should consider integrating their data.'),

        ('Fluxx',
         'Acquired: Oct 2021 by Metamorph Partners + ABS Capital Partners\n'
         'Previous Funding: $26.2M raised\n'
         'Business Model: Grantmaker-side SaaS (NOT grantseeker)\n'
         'Data: Grantelligence analytics (7,000+ visualizations)\n'
         'Key Strength: Serves funder/grantor side — enterprise analytics\n'
         'Key Weakness: Not a competitor — serves the other side of the market\n'
         'Opportunity: Partnership — if Fluxx funders could see GrantAQ readiness scores, funders send traffic to us.'),
    ]

    for name, profile in competitors:
        doc.add_heading(name, level=3)
        doc.add_paragraph(profile)

    # ── WHERE COMPETITORS GET DATA ──
    doc.add_heading('Where Every Competitor Gets Their Data', level=2)

    data_table = doc.add_table(rows=9, cols=4)
    data_table.style = 'Light Grid Accent 1'
    for i, h in enumerate(['Platform', 'Primary Data Sources', 'Update Frequency', 'Data Quality']):
        data_table.rows[0].cells[i].text = h

    data_rows = [
        ['Instrumentl', 'Grants.gov API + IRS 990s + state DBs + 27K manual-curated opps', 'Daily (federal), weekly (state), ongoing (manual)', 'High — human-verified'],
        ['GrantWatch', '40+ human researchers manually curating from govt sites + foundations', 'Daily updates by research team', 'Very high — all human-verified'],
        ['GrantStation', '25yr proprietary DB + direct funder relationships + manual research', 'Ongoing', 'High — relationship-based'],
        ['Grantable', 'IRS 990 filings (ProPublica API) + Grants.gov', '990s: 6+ month lag; Federal: real-time', 'Medium — 990 data is historical'],
        ['Granted AI', 'Federal + state + foundation DBs (133K foundations)', 'Not disclosed', 'Medium'],
        ['Candid', 'IRS 990s/990-PFs + CA AG Registry + direct submissions from 325K grantmakers', 'Processes ~3M grants/year', 'Highest — authoritative source'],
        ['GrantHub', 'Partnership with Candid (28M+ grants)', 'Via Candid partnership', 'High — Candid-powered'],
        ['GrantAQ', 'AI crawler (471 sources) + Grants.gov + state DBs + GPT-4o extraction', 'Daily/weekly/biweekly/monthly per source', 'Medium-High — AI-extracted, improving'],
    ]
    for r, row_data in enumerate(data_rows, 1):
        for c, val in enumerate(row_data):
            data_table.rows[r].cells[c].text = val
    doc.add_paragraph('')

    # ── FEATURE COMPARISON MATRIX ──
    doc.add_heading('Feature Comparison: GrantAQ vs Everyone', level=2)

    feat_table = doc.add_table(rows=16, cols=8)
    feat_table.style = 'Light Grid Accent 1'
    feat_headers = ['Feature', 'GrantAQ', 'Instrumentl', 'GrantWatch', 'Grantable', 'Granted AI', 'Candid', 'Consultants']
    for i, h in enumerate(feat_headers):
        feat_table.rows[0].cells[i].text = h

    features = [
        ['Grant Discovery/Matching', 'YES (AI)', 'YES (AI)', 'Search only', 'YES (990)', 'YES', 'Research', 'Manual'],
        ['Free Eligibility Check', 'YES', 'No', 'No', 'No', 'No', 'No', 'No'],
        ['Readiness Diagnostic', 'YES (AI, 30s)', 'No', 'No', 'No', 'No', 'No', '$2,500+'],
        ['COSO Controls Assessment', 'YES', 'No', 'No', 'No', 'No', 'No', 'Rare'],
        ['Site-Visit Simulation', 'YES', 'No', 'No', 'No', 'No', 'No', 'Rare'],
        ['If Not Ready Remediation', 'YES (detailed)', 'No', 'No', 'No', 'No', 'No', 'Manual'],
        ['AI Grant Writing', 'YES', 'No', 'No', 'YES', 'YES', 'No', 'No'],
        ['Expert Writing', 'YES', 'No', 'No', 'No', 'No', 'No', 'YES ($2K-$5K)'],
        ['Success Fee Model', 'YES (3-10%)', 'No', 'No', 'No', 'Guarantee', 'No', 'Rare'],
        ['Email Automation (22)', 'YES', 'No', 'No', 'Basic', 'No', 'No', 'Manual'],
        ['Lead Scoring + Alerts', 'YES', 'No', 'No', 'No', 'No', 'No', 'CRM needed'],
        ['Service Tiers (1-3)', 'YES', 'No', 'No', 'No', 'No', 'No', 'Custom'],
        ['Restructuring Options', 'YES', 'No', 'No', 'No', 'No', 'No', 'Custom'],
        ['Bootstrap Path', 'YES', 'No', 'No', 'No', 'No', 'No', 'No'],
        ['Pipeline Management', 'YES', 'YES', 'No', 'No', 'No', 'No', 'Spreadsheets'],
    ]
    for r, row_data in enumerate(features, 1):
        for c, val in enumerate(row_data):
            feat_table.rows[r].cells[c].text = val
    doc.add_paragraph('')

    # ── COMPETITIVE ADVANTAGES ──
    doc.add_heading('What We Have That Nobody Else Has', level=2)
    advantages = [
        ('Free Eligibility Check (no login)', 'Zero competitors offer this. $0.15 cost, captures leads. Our #1 growth lever.'),
        ('10-Step Readiness Diagnostic', 'Consultants charge $2,500-$25,000. We do it in 30 seconds with AI. Nobody else has automated this.'),
        ('"If Not Ready" Blocks', 'Every gap: what to do, specific URL, who, time, cost, dependencies, risk, which tier covers it.'),
        ('Path to Yes', 'Every "not eligible" category gets a fix path. Turns rejection into a sale.'),
        ('Bootstrap Path', 'Free/low-cost actions for tight-budget orgs. Builds goodwill, creates re-engagement loop.'),
        ('Restructuring for Ineligible', 'Standalone $1,497 product for "Not Eligible" verdicts. No competitor has this.'),
        ('22-Email Automated Nurture', 'Verdict-specific, data-personalized. Competitors have basic drip at best.'),
        ('Combined Discovery + Diagnostics + Writing + Services', 'Instrumentl does discovery. Grantable does writing. Consultants do readiness. We do ALL.'),
        ('Parallel AI Pipeline', '4 GPT-4o-mini + 1 GPT-4o synthesis = 18 seconds. Competitors take days/weeks.'),
        ('Full Confidence (Success Fee)', 'Win-or-pay-nothing model. Only Granted AI has a guarantee, but not success-fee based.'),
    ]
    for title, desc in advantages:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        p.add_run(desc)
    doc.add_paragraph('')

    # ── EVERY WAY TO MAKE GRANTAQ MORE SUCCESSFUL ──
    doc.add_heading('Every Way to Make GrantAQ More Successful', level=2)

    strategies = [
        ('ACQUISITION — How to Get Customers', [
            ('1. SEO for "am I eligible for grants" keywords', 'The /check page is your SEO asset. Target: "am I eligible for grants", "grant eligibility check", "small business grant readiness", "can my LLC get a grant". High-intent, low-competition. Every organic visitor becomes a lead at $0 CAC.'),
            ('2. LinkedIn content → /check funnel', 'Post 3x/week about grant readiness. Every post ends with "Check your eligibility free at grantaq.com/check". Target: founders, nonprofit EDs, small business owners. 1,000 impressions → 50 clicks → 25 checks → 1-2 signups.'),
            ('3. CPA/Accountant referral program', 'CPAs and bookkeepers serve your exact audience. Offer 10-20% referral fee on Tier 2/3. One active CPA partner = 5-10 qualified leads/month.'),
            ('4. SCORE mentor partnership', 'SCORE mentors advise 300K+ entrepreneurs/year. Many clients ask about grants. Become their recommended tool. One SCORE chapter = 50+ referrals/year.'),
            ('5. Webinar: "Is Your Organization Grant-Ready?"', 'Monthly webinar teaching grant readiness basics. End with free check. 100 registrations → 40 attend → 20 run check → 3-5 convert to Tier 1/2.'),
            ('6. State-specific landing pages', 'Create /grants/florida, /grants/texas, etc. Each page: state grant data + eligibility check CTA. Captures long-tail SEO.'),
            ('7. YouTube content on grant readiness', 'Videos: "5 Things That Block You From Getting Grants", "How to Register on SAM.gov", "Is My LLC Eligible for Grants?" End each with /check CTA.'),
            ('8. Podcast guesting', 'Target: small business, nonprofit, entrepreneur podcasts. Topic: "Why 90% of first-time grant applicants fail (and how to fix it)". Provide /check as resource.'),
            ('9. Facebook/IG ads targeting nonprofit pages', 'Audience: admins of nonprofit Facebook pages, followers of grant-related accounts. Ad: "Check your grant eligibility in 60 seconds — free."'),
            ('10. Cold outreach to SBA counselors', '900+ SBA offices nationwide. Each has counselors advising businesses on funding. Position GrantAQ as their grant readiness tool.'),
        ]),
        ('CONVERSION — How to Turn Leads Into Revenue', [
            ('11. Free diagnostic as the sales pitch', 'Every gap in the diagnostic points to a service tier. The diagnostic IS the sales meeting.'),
            ('12. Verdict-specific email sequences', '4 different email sequences based on verdict. Eligible Now gets urgency. Not Eligible gets restructuring options.'),
            ('13. Calendly booking in every email', 'Every report email and diagnostic has a "Book a 15-min call" link. 15 minutes is enough to close Tier 1/2.'),
            ('14. Stripe payment links in post-call email', 'Send within 1 hour of discovery call. Friction kills conversion — make it one click.'),
            ('15. Annual discount on subscriptions', '2 months free on annual plans. Locks in revenue and reduces churn.'),
            ('16. "Starter Grant Package" for $299', 'Write and submit 3-5 starter grant applications. Low risk, builds track record, leads to bigger engagements.'),
        ]),
        ('RETENTION — How to Keep Them', [
            ('17. Weekly grant matches email digest', 'Already built. Keeps users engaged even when not logging in.'),
            ('18. Pipeline deadline reminders', 'Automated alerts when grant deadlines approach. Creates urgency to stay subscribed.'),
            ('19. Annual re-diagnostic reminder', '12-month automated email with 50% returning client discount. Recurring revenue from one-time sales.'),
            ('20. Post-engagement 30-day check-in', 'Automated email asking how implementation is going. Upsell path to next tier.'),
            ('21. Success celebration emails', 'When a user marks a grant as "won" in pipeline, send congratulations + case study request + upsell to writing for next grant.'),
        ]),
        ('EXPANSION — How to Grow Revenue Per Customer', [
            ('22. Tier upgrade path', 'Free → Seeker → Strategist → Applicant → Organization. Each tier unlocks features that create demand for the next.'),
            ('23. Service tier ladder', 'Diagnostic → Tier 1 → Tier 2 → Tier 3 → Writing → Full Confidence. Natural progression.'),
            ('24. Multi-org accounts for consultants', 'Consultants managing 5-10 clients at $349/mo = $3,490 MRR from one customer.'),
            ('25. White-label for grant consulting firms', 'Let firms brand GrantAQ as their own platform. Enterprise deal at $1,000+/mo per firm.'),
            ('26. API access for integrators', 'Charge for API access to eligibility check + matching engine. Other platforms integrate with you.'),
        ]),
        ('MOAT — How to Make It Hard to Leave', [
            ('27. QuickBooks/Xero integration', 'Pull financial data directly for diagnostics. Once connected, switching costs increase dramatically.'),
            ('28. Grant-Ready Certification badge', 'After Tier 2/3, give orgs a "Grant-Ready Certified by GrantAQ" badge. Free marketing + social proof.'),
            ('29. Funder partnerships', 'Offer funders a dashboard showing applicant readiness scores. If funders send you traffic, you win.'),
            ('30. Compliance calendar integration', 'SAM.gov renewal, 990 filing, annual report reminders. Becomes their compliance system.'),
            ('31. Grant portfolio tracking', 'Track all active grants, reporting deadlines, award amounts. Becomes their grant CRM.'),
            ('32. Community/forum for grant seekers', 'Members-only community where grant seekers share tips, ask questions, celebrate wins. Network effect.'),
        ]),
        ('DATA — How to Build a Better Product', [
            ('33. Candid/FDO data partnership', 'Candid has 325K+ grantmakers and 2.4M+ grants. Integrating their data would instantly 50x your database.'),
            ('34. IRS 990 mining', 'Grantable mines 130K foundations from 990 data. You should too — it\'s free, public data.'),
            ('35. State database aggregation', 'Biggest gap in the market. No platform has comprehensive state/local grants. First-mover advantage.'),
            ('36. User feedback loop', 'When users mark grants as "applied" or "won", feed that back into the matching algorithm. Gets smarter over time.'),
            ('37. Funder behavior tracking', 'Track which funders are actively giving, changing priorities, or going dormant. Real-time funder intelligence.'),
        ]),
    ]

    for section_title, items in strategies:
        doc.add_heading(section_title, level=3)
        for title, desc in items:
            p = doc.add_paragraph()
            run = p.add_run(f'{title}: ')
            run.bold = True
            p.add_run(desc)
        doc.add_paragraph('')

    # ── REVENUE PROJECTIONS ──
    doc.add_heading('Revenue Projections', level=2)
    rev_table = doc.add_table(rows=7, cols=5)
    rev_table.style = 'Light Grid Accent 1'
    for i, h in enumerate(['Revenue Stream', 'Month 3', 'Month 6', 'Month 12', 'Assumptions']):
        rev_table.rows[0].cells[i].text = h
    projections = [
        ['Subscriptions (MRR)', '$2,500', '$12,000', '$49,500', '50 → 200 → 500 subscribers at $50-$99 avg'],
        ['Service Tiers', '$4,000', '$20,000', '$40,000', '2 → 10 → 20 sales/month at $2,000 avg'],
        ['Writing (per-grant)', '$1,500', '$8,000', '$20,000', '5 → 20 → 50 purchases/month at $400 avg'],
        ['Success Fees', '$0', '$2,000', '$10,000', 'Starts Month 6 as first clients win awards'],
        ['TOTAL MONTHLY', '$8,000', '$42,000', '$119,500', ''],
        ['TOTAL ANNUAL RUN RATE', '$96,000', '$504,000', '$1,434,000', ''],
    ]
    for r, row_data in enumerate(projections, 1):
        for c, val in enumerate(row_data):
            rev_table.rows[r].cells[c].text = val

    doc.add_paragraph('')
    doc.add_paragraph('— End of Report —')

    doc.save('/Users/poweredbyexcellence/grantiq/docs/pricing/GrantAQ_Pricing_Strategy.docx')
    print('Word doc saved')


# ============================================================================
# EXCEL SPREADSHEET
# ============================================================================

def create_excel():
    wb = openpyxl.Workbook()

    hf = Font(bold=True, color='FFFFFF', size=11)
    hfill = PatternFill(start_color='0D9488', end_color='0D9488', fill_type='solid')
    bf = Font(bold=True, size=11)
    wrap = Alignment(wrap_text=True, vertical='top')
    border = Border(
        left=Side(style='thin', color='E7E5E4'), right=Side(style='thin', color='E7E5E4'),
        top=Side(style='thin', color='E7E5E4'), bottom=Side(style='thin', color='E7E5E4'),
    )

    def style_header(ws, row=1, cols=10):
        for c in range(1, cols + 1):
            cell = ws.cell(row=row, column=c)
            cell.font = hf
            cell.fill = hfill
            cell.alignment = Alignment(wrap_text=True, vertical='center')
            cell.border = border

    def style_rows(ws, sr, er, cols):
        for r in range(sr, er + 1):
            for c in range(1, cols + 1):
                cell = ws.cell(row=r, column=c)
                cell.alignment = wrap
                cell.border = border

    # ── Sheet 1: Subscription Pricing ──
    ws1 = wb.active
    ws1.title = 'Subscriptions'
    h1 = ['Tier', 'Name', 'Monthly', 'Annual', 'Monthly (Ann)', 'Features', 'Target', 'Competitor']
    for c, h in enumerate(h1, 1): ws1.cell(row=1, column=c, value=h)
    style_header(ws1, cols=8)
    t1 = [
        ['free', 'Explorer', '$0', 'Forever', '$0', '1 match/mo, federal only, readiness, free check', 'Lead capture', 'No competitor free tier with AI'],
        ['starter', 'Seeker', '$39', '$390', '$32.50', 'Full library, 10 pipeline, calendar, 5 docs', 'Small orgs', 'Undercuts GrantWatch $49, Grantable $50'],
        ['pro', 'Strategist', '$99', '$950', '$79', 'Unlimited scorecard, vault, 30 chats, analytics', 'Active seekers', 'Between Grantable Pro $150, GrantWatch $49'],
        ['growth', 'Applicant', '$179', '$1,710', '$143', 'AI writing, expert review, unlimited pipeline', 'Writing apps', 'Instrumentl $299 NO writing'],
        ['enterprise', 'Organization', '$349', '$3,350', '$279', 'All + writer, team, API, CSM', 'Multi-user', 'Instrumentl $899 less features'],
    ]
    for r, d in enumerate(t1, 2):
        for c, v in enumerate(d, 1): ws1.cell(row=r, column=c, value=v)
    style_rows(ws1, 2, 6, 8)
    for c in range(1, 9): ws1.column_dimensions[get_column_letter(c)].width = 18 if c < 6 else 40

    # ── Sheet 2: Service Tiers ──
    ws2 = wb.create_sheet('Services')
    h2 = ['Service', 'Price', 'Delivery', 'Included', 'Why', 'Competitor', 'Competitor Price']
    for c, h in enumerate(h2, 1): ws2.cell(row=1, column=c, value=h)
    style_header(ws2, cols=7)
    s2 = [
        ['Eligibility Status', 'FREE', 'Instant', 'Verdict, categories, blockers, quick wins', 'Lead magnet $0.15 cost', 'None exist', 'N/A'],
        ['Readiness Diagnostic', 'FREE (acct)', '30 sec', '10-step diagnostic', 'Drives signups', 'Consultant', '$2,500-$25,000'],
        ['Tier 1 Review', '$497', '5-7 days', 'Diagnostic + 45min call', '80% below consultants', 'Consultant', '$2,500-$5,000'],
        ['Tier 2 Roadmap', '$1,997', '2-3 weeks', 'Playbook + templates + 2 calls', '12x ROI on $25K grant', 'Consultant', '$5,000-$15,000'],
        ['Tier 3 Accelerator', '$4,997', '60-120 days', 'Done-for-you + first app', '50-75% below consulting', 'Consultant', '$10,000-$25,000+'],
        ['Restructuring', '$1,497', '2-4 weeks', 'Structural analysis + alt capital', 'Unique product', 'None exist', 'N/A'],
    ]
    for r, d in enumerate(s2, 2):
        for c, v in enumerate(d, 1): ws2.cell(row=r, column=c, value=v)
    style_rows(ws2, 2, 7, 7)
    for c in range(1, 8): ws2.column_dimensions[get_column_letter(c)].width = 22

    # ── Sheet 3: Deep Competitor Intel ──
    ws3 = wb.create_sheet('Competitors Deep')
    h3 = ['Company', 'Pricing', 'Revenue/ARR', 'Funding', 'Customers', 'Team Size', 'Database', 'Data Sources', 'Strength', 'Weakness', 'We Beat Them On']
    for c, h in enumerate(h3, 1): ws3.cell(row=1, column=c, value=h)
    style_header(ws3, cols=11)
    c3 = [
        ['Instrumentl', '$299-$899/mo', 'Private', '$55M+', '4,500+', 'Expanding', '450K funders, 27K opps', 'Grants.gov + 990s + manual', 'Brand, funding, customers', 'No readiness, no writing, expensive', 'Diagnostics + writing + services'],
        ['Submittable', 'Custom', '$66.6M (2024)', '$74.6M', '11,000', '230', 'Grant mgmt (not discovery)', 'Platform submissions', 'Revenue leader', 'Not discovery — different market', 'Different segment entirely'],
        ['GrantWatch', '$49/mo', 'Private', 'Bootstrapped', '350K+ users', '40+ researchers', '10K-11.5K grants', '40+ human researchers', 'Largest user base, quality', 'No AI, no matching, just directory', 'AI matching + diagnostics + writing'],
        ['GrantStation', '$699/yr', 'Private', 'PE (Elios)', 'Not disclosed', 'Not disclosed', '150K profiles, 15K grants', '25yr DB + funder relationships', '25yr history, credibility', 'Dated UI, no AI, expensive', 'AI + modern UX + lower price'],
        ['Grantable', '$22-$68/mo', 'Private', '$100K grants', '13,000+', 'Small', '130K foundations', 'IRS 990 + Grants.gov', 'AI coworker, affordable', 'No readiness, no expert writing', 'Diagnostics + expert writing + services'],
        ['Granted AI', '$57/mo', 'Private', 'Not disclosed', 'Not disclosed', 'Small', '133K foundations', 'Federal + state + foundation', 'Guarantee, unlimited drafts', 'Small DB, no readiness', 'Matching + diagnostics + services'],
        ['Candid/FDO', '$35-$127/mo', 'Nonprofit', 'Nonprofit', '200+ integrations', 'Large', '325K grantmakers, 2.4M grants', '990s + PFs + direct submissions', 'THE data authority', 'Research only, no app help', 'Full platform vs research tool'],
        ['GrantAQ', '$0-$349/mo', 'Early stage', 'Self-funded', 'Growing', 'Small', '4,300+ grants', 'AI crawler + Grants.gov + state', 'Only all-in-one platform', 'Smaller DB (growing)', 'EVERYTHING combined'],
    ]
    for r, d in enumerate(c3, 2):
        for c, v in enumerate(d, 1): ws3.cell(row=r, column=c, value=v)
    style_rows(ws3, 2, 9, 11)
    for c in range(1, 12): ws3.column_dimensions[get_column_letter(c)].width = 25

    # ── Sheet 4: Feature Matrix ──
    ws4 = wb.create_sheet('Feature Matrix')
    h4 = ['Feature', 'GrantAQ', 'Instrumentl', 'GrantWatch', 'Grantable', 'Granted AI', 'Candid', 'Consultants']
    for c, h in enumerate(h4, 1): ws4.cell(row=1, column=c, value=h)
    style_header(ws4, cols=8)
    f4 = [
        ['Grant Discovery/Matching', 'YES (AI)', 'YES (AI)', 'Search only', 'YES (990)', 'YES', 'Research', 'Manual'],
        ['Free Eligibility Check', 'YES', 'No', 'No', 'No', 'No', 'No', 'No'],
        ['Readiness Diagnostic', 'YES (AI, 30s)', 'No', 'No', 'No', 'No', 'No', '$2,500+'],
        ['COSO Controls Assessment', 'YES', 'No', 'No', 'No', 'No', 'No', 'Rare'],
        ['Site-Visit Simulation', 'YES', 'No', 'No', 'No', 'No', 'No', 'Rare'],
        ['Remediation Blocks', 'YES (detailed)', 'No', 'No', 'No', 'No', 'No', 'Manual'],
        ['AI Grant Writing', 'YES', 'No', 'No', 'YES', 'YES', 'No', 'No'],
        ['Expert Writing Option', 'YES', 'No', 'No', 'No', 'No', 'No', '$2K-$5K'],
        ['Success Fee Model', 'YES (3-10%)', 'No', 'No', 'No', 'Guarantee', 'No', 'Rare'],
        ['22-Email Automation', 'YES', 'No', 'No', 'Basic', 'No', 'No', 'Manual'],
        ['Lead Scoring', 'YES', 'No', 'No', 'No', 'No', 'No', 'CRM needed'],
        ['Service Tiers 1-3', 'YES', 'No', 'No', 'No', 'No', 'No', 'Custom'],
        ['Restructuring Options', 'YES', 'No', 'No', 'No', 'No', 'No', 'Custom'],
        ['Bootstrap Path', 'YES', 'No', 'No', 'No', 'No', 'No', 'No'],
        ['Pipeline Mgmt', 'YES', 'YES', 'No', 'No', 'No', 'No', 'Spreadsheets'],
    ]
    for r, d in enumerate(f4, 2):
        for c, v in enumerate(d, 1):
            cell = ws4.cell(row=r, column=c, value=v)
            if v == 'YES' or v.startswith('YES'):
                cell.font = Font(bold=True, color='0D9488')
            elif v == 'No':
                cell.font = Font(color='999999')
    style_rows(ws4, 2, 16, 8)
    for c in range(1, 9): ws4.column_dimensions[get_column_letter(c)].width = 20

    # ── Sheet 5: Growth Strategies ──
    ws5 = wb.create_sheet('Growth Strategies')
    h5 = ['#', 'Category', 'Strategy', 'Description', 'Impact', 'Effort', 'Priority']
    for c, h in enumerate(h5, 1): ws5.cell(row=1, column=c, value=h)
    style_header(ws5, cols=7)
    g5 = [
        [1, 'Acquisition', 'SEO for eligibility keywords', '/check = landing page for "am I eligible for grants"', 'High', 'Medium', 'P1'],
        [2, 'Acquisition', 'LinkedIn → /check', 'Post 3x/week, CTA to free check', 'Medium', 'Low', 'P1'],
        [3, 'Acquisition', 'CPA referral program', '10-20% fee on Tier 2/3', 'High', 'Medium', 'P1'],
        [4, 'Acquisition', 'SCORE partnership', 'Recommended tool for 300K+ entrepreneurs', 'High', 'High', 'P2'],
        [5, 'Acquisition', 'Monthly webinar', '"Is Your Org Grant-Ready?" → free check', 'Medium', 'Medium', 'P2'],
        [6, 'Acquisition', 'State landing pages', '/grants/florida etc', 'High', 'Medium', 'P2'],
        [7, 'Acquisition', 'YouTube content', 'Grant readiness videos → /check CTA', 'High', 'Medium', 'P2'],
        [8, 'Acquisition', 'Podcast guesting', 'Entrepreneur/nonprofit podcasts', 'Medium', 'Low', 'P2'],
        [9, 'Acquisition', 'FB/IG ads to nonprofits', 'Target nonprofit page admins', 'Medium', 'Medium', 'P3'],
        [10, 'Acquisition', 'SBA counselor outreach', '900+ offices, position as their grant tool', 'High', 'High', 'P3'],
        [11, 'Conversion', 'Free diagnostic as sales pitch', 'Every gap → service tier', 'Very High', 'Done', 'P1'],
        [12, 'Conversion', 'Verdict-specific emails', '4 email paths by verdict', 'High', 'Done', 'P1'],
        [13, 'Conversion', 'Calendly in every email', '15-min call link in report + emails', 'High', 'Low', 'P1'],
        [14, 'Conversion', 'Stripe links post-call', 'One-click payment in follow-up email', 'High', 'Low', 'P1'],
        [15, 'Conversion', 'Annual discount', '2 months free on annual plans', 'Medium', 'Low', 'P1'],
        [16, 'Conversion', 'Starter grant package $299', '3-5 starter apps submitted', 'Medium', 'Medium', 'P2'],
        [17, 'Retention', 'Weekly digest emails', 'Grant matches + deadlines', 'High', 'Done', 'P1'],
        [18, 'Retention', 'Pipeline deadline alerts', 'Auto reminders for approaching deadlines', 'Medium', 'Done', 'P1'],
        [19, 'Retention', 'Annual re-diagnostic', '12-month reminder + 50% discount', 'Medium', 'Done', 'P1'],
        [20, 'Retention', '30-day post-engagement', 'Check-in email → upsell next tier', 'Medium', 'Done', 'P1'],
        [21, 'Expansion', 'Tier upgrade path', 'Free → Seeker → Strategist → Applicant → Org', 'Very High', 'Done', 'P1'],
        [22, 'Expansion', 'Service tier ladder', 'Diag → T1 → T2 → T3 → Writing', 'Very High', 'Done', 'P1'],
        [23, 'Expansion', 'Multi-org for consultants', '5-10 clients at $349/mo = $3,490 MRR', 'High', 'Medium', 'P2'],
        [24, 'Expansion', 'White-label for firms', 'Brand as their own platform $1K+/mo', 'Very High', 'High', 'P3'],
        [25, 'Expansion', 'API access', 'Charge for eligibility + matching API', 'High', 'High', 'P3'],
        [26, 'Moat', 'QuickBooks/Xero integration', 'Pull financial data → switching cost', 'Very High', 'High', 'P3'],
        [27, 'Moat', 'Grant-Ready Certification', 'Badge for Tier 2/3 → free marketing', 'Medium', 'Low', 'P2'],
        [28, 'Moat', 'Funder partnerships', 'Readiness dashboard for funders', 'Very High', 'High', 'P3'],
        [29, 'Data', 'Candid data partnership', '325K grantmakers → 50x your DB', 'Very High', 'High', 'P2'],
        [30, 'Data', 'IRS 990 mining', '130K foundations free data', 'High', 'Medium', 'P2'],
        [31, 'Data', 'State DB aggregation', 'First-mover on state/local grants', 'Very High', 'High', 'P2'],
    ]
    for r, d in enumerate(g5, 2):
        for c, v in enumerate(d, 1): ws5.cell(row=r, column=c, value=v)
    style_rows(ws5, 2, 32, 7)
    ws5.column_dimensions['A'].width = 5
    ws5.column_dimensions['B'].width = 14
    ws5.column_dimensions['C'].width = 28
    ws5.column_dimensions['D'].width = 50
    ws5.column_dimensions['E'].width = 12
    ws5.column_dimensions['F'].width = 10
    ws5.column_dimensions['G'].width = 8

    # ── Sheet 6: Revenue Projections ──
    ws6 = wb.create_sheet('Revenue')
    h6 = ['Stream', 'Month 3', 'Month 6', 'Month 12', 'Assumptions']
    for c, h in enumerate(h6, 1): ws6.cell(row=1, column=c, value=h)
    style_header(ws6, cols=5)
    p6 = [
        ['Subscriptions', '$2,500', '$12,000', '$49,500', '50→200→500 subs at $50-$99 avg'],
        ['Service Tiers', '$4,000', '$20,000', '$40,000', '2→10→20 sales/mo at $2K avg'],
        ['Writing', '$1,500', '$8,000', '$20,000', '5→20→50 purchases at $400 avg'],
        ['Success Fees', '$0', '$2,000', '$10,000', 'Starts M6 as clients win'],
        ['TOTAL MONTHLY', '$8,000', '$42,000', '$119,500', ''],
        ['ANNUAL RUN RATE', '$96,000', '$504,000', '$1,434,000', ''],
    ]
    for r, d in enumerate(p6, 2):
        for c, v in enumerate(d, 1):
            cell = ws6.cell(row=r, column=c, value=v)
            if r >= 6: cell.font = bf
    style_rows(ws6, 2, 7, 5)
    for c in range(1, 6): ws6.column_dimensions[get_column_letter(c)].width = 22

    wb.save('/Users/poweredbyexcellence/grantiq/docs/pricing/GrantAQ_Pricing_Analysis.xlsx')
    print('Excel saved')


if __name__ == '__main__':
    create_word_doc()
    create_excel()
    print('Done!')
